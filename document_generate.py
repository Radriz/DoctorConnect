import os
from io import BytesIO

from docxtpl import DocxTemplate
from docx import Document
from docx.shared import Pt
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def generate_plan(
        stages,
        data,
        total_price,
        output_path = os.getcwd() + "/templates/file/output.docx",
        template_path = os.getcwd() + "/templates/file/plan_word.docx",
        add_template_path = os.getcwd()  + "/templates/file/add_plan_word.docx",
):
    def fill_docx_template(doc, replacements):
        for para in doc.paragraphs:
            for run in para.runs:
                for key, value in replacements.items():
                    if key in run.text:
                        run.text = run.text.replace(key, value)
        return doc
    def set_cell_border(cell, **kwargs):
        tc = cell._element
        tcPr = tc.get_or_add_tcPr()

        for edge in ('top', 'bottom', 'start', 'end'):
            edge_data = kwargs.get(edge)
            if edge_data:
                tag = f'w:{edge}'
                element = tcPr.find(qn(tag))
                if element is None:
                    element = OxmlElement(tag)
                    tcPr.append(element)

                for key in ["sz", "val", "color", "space"]:
                    if key in edge_data:
                        element.set(qn(f'w:{key}'), str(edge_data[key]))

    def create_table(doc, items):
        table = doc.add_table(rows=1, cols=5)
        hdr_cells = table.rows[0].cells
        headers = ['№', 'Услуга', 'Цена за ед.', 'Кол-во', 'Всего']

        for i, header in enumerate(headers):
            cell = hdr_cells[i]
            run = cell.paragraphs[0].add_run(header)
            run.bold = True
            run.font.size = Pt(14)

            set_cell_border(cell, top={"sz": 10, "val": "single", "color": "000000"},
                            bottom={"sz": 10, "val": "single", "color": "000000"},
                            start={"sz": 10, "val": "single", "color": "000000"},
                            end={"sz": 10, "val": "single", "color": "000000"})

        for item in items:
            row_cells = table.add_row().cells
            row_cells[0].text = str(item['no'])
            row_cells[1].text = item['service'] + "  "
            row_cells[2].text = f"{item['price_per_unit']} руб."
            row_cells[3].text = str(item['quantity'])
            row_cells[4].text = f"{item['total']} руб."
            for cell in row_cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(14)

                set_cell_border(cell, top={"sz": 10, "val": "single", "color": "000000"},
                                bottom={"sz": 10, "val": "single", "color": "000000"},
                                start={"sz": 10, "val": "single", "color": "000000"},
                                end={"sz": 10, "val": "single", "color": "000000"})

    # Генерация основного документа
    doc = Document(template_path)
    doc = fill_docx_template(doc, data)

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    document = Document(buffer)

    for stage in stages:
        title = document.add_paragraph()
        ts = title.add_run(f"-" * 132)
        ts = title.add_run(f"Этап №{stage['number']} – {stage['stage']}")
        ts.font.size = Pt(14)
        title = document.add_paragraph()
        for stage_template in stage['templates'].values():
            ts = title.add_run(f"{stage_template['name']} — {stage['tooth']}" + " " + ("зуб" if ',' not in stage['tooth'] else "зубы"))
            ts.font.size = Pt(14)
            create_table(document, stage_template['items'])
            title = document.add_paragraph()
            title.paragraph_format.space_before = Pt(13)
        ts = title.add_run(f"Общая стоимость за этап {stage['stage'].lower()}: {stage['total_price']} руб.")
        ts.font.size = Pt(14)

    title = document.add_paragraph()
    ts = title.add_run(f"-" * 132)
    title = document.add_paragraph()
    title.paragraph_format.space_after = Pt(20)
    ts = title.add_run("ОБЩАЯ СТОИМОСТЬ ЗА ПЛАН ЛЕЧЕНИЯ: ")
    ts.bold = True
    ts.font.size = Pt(14)
    ts = title.add_run(f"{total_price} руб.")
    ts.bold = True
    ts.font.size = Pt(18)

    # Загружаем дополнительный шаблон и добавляем его к основному документу
    doc_end = Document(add_template_path)

    for element in doc_end.element.body:
        document.element.body.append(element)

    # Сохраняем результирующий документ
    document.save(output_path)

